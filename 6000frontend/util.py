import io
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def build_resume_html(
    template_name,
    accent_color,
    name,
    email,
    summary,
    skills,
    experience,
    education,
    generated_resume
):
    """
    Build a styled HTML resume preview for Streamlit display.
    """
    skills_html = "".join(
        [f"<span class='skill-pill' style='border-color:{accent_color}; color:{accent_color};'>{skill}</span>" for skill in skills]
    )

    if template_name == "Modern":
        return f"""
        <div class="resume-template modern-template">
            <div class="resume-banner" style="background:{accent_color};"></div>

            <div class="resume-header modern-header">
                <div>
                    <h1>{name}</h1>
                    <p>{email}</p>
                </div>
            </div>

            <div class="resume-grid">
                <div class="left-panel">
                    <h3 style="color:{accent_color};">Profile</h3>
                    <p>{summary}</p>

                    <h3 style="color:{accent_color};">Skills</h3>
                    <div class="skills-wrap">{skills_html}</div>

                    <h3 style="color:{accent_color};">Education</h3>
                    <p>{education}</p>
                </div>

                <div class="right-panel">
                    <h3 style="color:{accent_color};">Experience</h3>
                    <p>{experience.replace(chr(10), '<br>')}</p>

                    <h3 style="color:{accent_color};">AI-Enhanced Resume</h3>
                    <p>{generated_resume.replace(chr(10), '<br>')}</p>
                </div>
            </div>
        </div>
        """

    if template_name == "Minimal":
        return f"""
        <div class="resume-template minimal-template">
            <div class="resume-header">
                <h1>{name}</h1>
                <p>{email}</p>
            </div>

            <div class="resume-section">
                <h3 style="color:{accent_color};">Summary</h3>
                <p>{summary}</p>
            </div>

            <div class="resume-section">
                <h3 style="color:{accent_color};">Skills</h3>
                <div class="skills-wrap">{skills_html}</div>
            </div>

            <div class="resume-section">
                <h3 style="color:{accent_color};">Experience</h3>
                <p>{experience.replace(chr(10), '<br>')}</p>
            </div>

            <div class="resume-section">
                <h3 style="color:{accent_color};">Education</h3>
                <p>{education}</p>
            </div>

            <div class="resume-section">
                <h3 style="color:{accent_color};">Optimised Resume</h3>
                <p>{generated_resume.replace(chr(10), '<br>')}</p>
            </div>
        </div>
        """

    return f"""
    <div class="resume-template professional-template">
        <div class="resume-header" style="border-bottom: 3px solid {accent_color};">
            <h1>{name}</h1>
            <p>{email}</p>
        </div>

        <div class="resume-section">
            <h3 style="color:{accent_color};">Professional Summary</h3>
            <p>{summary}</p>
        </div>

        <div class="resume-section">
            <h3 style="color:{accent_color};">Key Skills</h3>
            <div class="skills-wrap">{skills_html}</div>
        </div>

        <div class="resume-section">
            <h3 style="color:{accent_color};">Experience</h3>
            <p>{experience.replace(chr(10), '<br>')}</p>
        </div>

        <div class="resume-section">
            <h3 style="color:{accent_color};">Education</h3>
            <p>{education}</p>
        </div>

        <div class="resume-section">
            <h3 style="color:{accent_color};">AI-Enhanced Resume</h3>
            <p>{generated_resume.replace(chr(10), '<br>')}</p>
        </div>
    </div>
    """


def build_resume_text(template_name, name, email, summary, skills, experience, education, generated_resume):
    """
    Build plain text version for TXT and PDF export.
    """
    skills_text = ", ".join(skills)

    return f"""
{name}
Email: {email}

PROFESSIONAL SUMMARY
{summary}

SKILLS
{skills_text}

EXPERIENCE
{experience}

EDUCATION
{education}

AI-ENHANCED RESUME
{generated_resume}
""".strip()


def generate_pdf_bytes(resume_text, title="AI Resume"):
    """
    Create PDF bytes for download.
    """
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    y = height - 50

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, title)
    y -= 30

    pdf.setFont("Helvetica", 10)

    for line in resume_text.split("\n"):
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = height - 50

        pdf.drawString(50, y, line[:100])
        y -= 15

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def generate_docx_bytes(name, email, summary, skills, experience, education, generated_resume):
    """
    Create editable DOCX bytes for users to download and edit later.
    """
    doc = Document()

    doc.add_heading(name, level=0)
    doc.add_paragraph(email)

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(skills))

    doc.add_heading("Experience", level=1)
    doc.add_paragraph(experience)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.add_heading("AI-Enhanced Resume Content", level=1)
    doc.add_paragraph(generated_resume)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()